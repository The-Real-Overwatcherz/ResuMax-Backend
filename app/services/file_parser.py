"""
ResuMax — File Parser Service
Extracts raw text from PDF, DOCX, and TXT resume files.
"""

import io
import re
import structlog

logger = structlog.get_logger(__name__)


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    import fitz  # PyMuPDF

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text_parts = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text_parts.append(page.get_text("text"))

    doc.close()
    raw_text = "\n".join(text_parts)
    logger.info("pdf_parsed", pages=len(text_parts), chars=len(raw_text))
    return raw_text


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Also extract from tables (some resumes use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    raw_text = "\n".join(text_parts)
    logger.info("docx_parsed", paragraphs=len(text_parts), chars=len(raw_text))
    return raw_text


def parse_txt(file_bytes: bytes) -> str:
    """Decode a plain text file with encoding fallback."""
    for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
        try:
            raw_text = file_bytes.decode(encoding)
            logger.info("txt_parsed", encoding=encoding, chars=len(raw_text))
            return raw_text
        except (UnicodeDecodeError, ValueError):
            continue

    # Last resort: decode with replacement chars
    raw_text = file_bytes.decode("utf-8", errors="replace")
    logger.warning("txt_parsed_with_errors", chars=len(raw_text))
    return raw_text


def clean_text(text: str) -> str:
    """
    Clean extracted resume text:
    - Normalize whitespace
    - Fix encoding artifacts
    - Remove excessive blank lines
    """
    # Fix common encoding artifacts
    text = text.replace("\u2019", "'")      # Smart single quote
    text = text.replace("\u2018", "'")
    text = text.replace("\u201c", '"')      # Smart double quotes
    text = text.replace("\u201d", '"')
    text = text.replace("\u2013", "-")      # En dash
    text = text.replace("\u2014", "-")      # Em dash
    text = text.replace("\u2022", "•")      # Bullet point
    text = text.replace("\uf0b7", "•")      # Another bullet
    text = text.replace("\xa0", " ")        # Non-breaking space

    # Normalize whitespace within lines
    text = re.sub(r"[ \t]+", " ", text)

    # Remove excessive blank lines (keep max 2)
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Strip leading/trailing whitespace per line
    lines = [line.strip() for line in text.split("\n")]
    text = "\n".join(lines)

    # Strip overall
    text = text.strip()

    return text


def parse_resume_file(file_bytes: bytes, filename: str) -> str:
    """
    Route to the correct parser based on file extension.
    Returns cleaned, normalized text.

    Raises:
        ValueError: If file type is not supported
    """
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        raw_text = parse_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        raw_text = parse_docx(file_bytes)
    elif filename_lower.endswith(".txt"):
        raw_text = parse_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {filename}. Accepted: .pdf, .docx, .txt")

    cleaned = clean_text(raw_text)

    if len(cleaned) < 50:
        raise ValueError("Could not extract meaningful text from the file. Please try a different format.")

    logger.info("resume_file_parsed", filename=filename, raw_chars=len(raw_text), clean_chars=len(cleaned))
    return cleaned
