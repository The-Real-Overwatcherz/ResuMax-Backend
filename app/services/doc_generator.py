"""
ResuMax — Document Generator Service
Generates a professional DOCX resume from the optimized pipeline output.
"""

import io
import structlog
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = structlog.get_logger(__name__)


def _add_section_heading(doc: Document, text: str):
    """Add a styled section heading with a bottom border."""
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    p.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    run.font.name = "Calibri"
    # Add a thin line under the heading
    from docx.oxml.ns import qn
    pBdr = p.paragraph_format.element.get_or_add_pPr()
    bdr = pBdr.makeelement(qn("w:pBdr"), {})
    bottom = bdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "4",
        qn("w:space"): "1",
        qn("w:color"): "333333",
    })
    bdr.append(bottom)
    pBdr.append(bdr)


def generate_optimized_resume(optimized_resume: dict, analysis_data: dict) -> io.BytesIO:
    """
    Generate a professional DOCX from the optimized resume data.
    Returns a BytesIO buffer containing the DOCX file.
    """
    doc = Document()

    # ── Page Margins ─────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # ── Default Style ────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)

    contact = optimized_resume.get("contact", {})
    
    # ── Header: Name ─────────────────────────────────────────
    name = contact.get("full_name", "Your Name")
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.space_after = Pt(2)
    name_run = name_para.add_run(name)
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    name_run.font.name = "Calibri"

    # ── Contact Info Line ────────────────────────────────────
    contact_parts = []
    if contact.get("email"):
        contact_parts.append(contact["email"])
    if contact.get("phone"):
        contact_parts.append(contact["phone"])
    if contact.get("location"):
        contact_parts.append(contact["location"])
    if contact.get("linkedin"):
        contact_parts.append(contact["linkedin"])
    if contact.get("github"):
        contact_parts.append(contact["github"])

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.space_after = Pt(6)
        contact_run = contact_para.add_run("  |  ".join(contact_parts))
        contact_run.font.size = Pt(9)
        contact_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        contact_run.font.name = "Calibri"

    # ── Summary ──────────────────────────────────────────────
    summary = optimized_resume.get("summary", "")
    if summary:
        _add_section_heading(doc, "Professional Summary")
        p = doc.add_paragraph()
        run = p.add_run(summary)
        run.font.size = Pt(10)
        run.font.name = "Calibri"

    # ── Experience ───────────────────────────────────────────
    experience = optimized_resume.get("experience", [])
    if experience:
        _add_section_heading(doc, "Experience")
        for exp in experience:
            # Title + Company line
            title_para = doc.add_paragraph()
            title_para.space_before = Pt(6)
            title_para.space_after = Pt(1)
            
            title_text = exp.get("title", "")
            company_text = exp.get("company", "")
            
            title_run = title_para.add_run(title_text)
            title_run.bold = True
            title_run.font.size = Pt(10.5)
            title_run.font.name = "Calibri"
            
            if company_text:
                sep_run = title_para.add_run(f"  —  {company_text}")
                sep_run.font.size = Pt(10)
                sep_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                sep_run.font.name = "Calibri"

            dates = exp.get("dates", "")
            if dates:
                dates_run = title_para.add_run(f"    ({dates})")
                dates_run.font.size = Pt(9)
                dates_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
                dates_run.font.name = "Calibri"

            # Bullet points
            for bullet in exp.get("bullets", []):
                bp = doc.add_paragraph(style="List Bullet")
                bp.space_before = Pt(1)
                bp.space_after = Pt(1)
                bp_run = bp.add_run(bullet)
                bp_run.font.size = Pt(10)
                bp_run.font.name = "Calibri"

    # ── Projects ─────────────────────────────────────────────
    projects = optimized_resume.get("projects", [])
    if projects:
        _add_section_heading(doc, "Projects")
        for proj in projects:
            if isinstance(proj, dict):
                proj_para = doc.add_paragraph()
                proj_para.space_before = Pt(4)
                proj_para.space_after = Pt(1)
                
                proj_name = proj.get("name", proj.get("title", ""))
                proj_run = proj_para.add_run(proj_name)
                proj_run.bold = True
                proj_run.font.size = Pt(10.5)
                proj_run.font.name = "Calibri"

                tech = proj.get("technologies", proj.get("tech_stack", []))
                if tech:
                    tech_str = ", ".join(tech) if isinstance(tech, list) else str(tech)
                    tech_run = proj_para.add_run(f"  [{tech_str}]")
                    tech_run.font.size = Pt(9)
                    tech_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                    tech_run.font.name = "Calibri"

                desc = proj.get("description", "")
                if desc:
                    desc_para = doc.add_paragraph()
                    desc_run = desc_para.add_run(desc)
                    desc_run.font.size = Pt(10)
                    desc_run.font.name = "Calibri"

                for bullet in proj.get("bullets", []):
                    bp = doc.add_paragraph(style="List Bullet")
                    bp_run = bp.add_run(bullet)
                    bp_run.font.size = Pt(10)
                    bp_run.font.name = "Calibri"
            elif isinstance(proj, str):
                bp = doc.add_paragraph(style="List Bullet")
                bp_run = bp.add_run(proj)
                bp_run.font.size = Pt(10)
                bp_run.font.name = "Calibri"

    # ── Education ────────────────────────────────────────────
    education = optimized_resume.get("education", [])
    if education:
        _add_section_heading(doc, "Education")
        for edu in education:
            if isinstance(edu, dict):
                edu_para = doc.add_paragraph()
                edu_para.space_before = Pt(4)
                edu_para.space_after = Pt(1)

                degree = edu.get("degree", "")
                institution = edu.get("institution", edu.get("school", ""))
                
                deg_run = edu_para.add_run(degree)
                deg_run.bold = True
                deg_run.font.size = Pt(10.5)
                deg_run.font.name = "Calibri"

                if institution:
                    inst_run = edu_para.add_run(f"  —  {institution}")
                    inst_run.font.size = Pt(10)
                    inst_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                    inst_run.font.name = "Calibri"

                dates = edu.get("dates", edu.get("graduation", ""))
                if dates:
                    date_run = edu_para.add_run(f"    ({dates})")
                    date_run.font.size = Pt(9)
                    date_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
                    date_run.font.name = "Calibri"
            elif isinstance(edu, str):
                p = doc.add_paragraph()
                p.add_run(edu).font.size = Pt(10)

    # ── Skills ───────────────────────────────────────────────
    skills = optimized_resume.get("skills", [])
    if skills:
        _add_section_heading(doc, "Skills")
        skills_text = ", ".join(skills) if isinstance(skills, list) else str(skills)
        p = doc.add_paragraph()
        run = p.add_run(skills_text)
        run.font.size = Pt(10)
        run.font.name = "Calibri"

    # ── Certifications ───────────────────────────────────────
    certs = optimized_resume.get("certifications", [])
    if certs:
        _add_section_heading(doc, "Certifications")
        for cert in certs:
            if isinstance(cert, dict):
                p = doc.add_paragraph(style="List Bullet")
                cert_text = cert.get("name", cert.get("title", str(cert)))
                run = p.add_run(cert_text)
                run.font.size = Pt(10)
                run.font.name = "Calibri"
            elif isinstance(cert, str):
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(cert)
                run.font.size = Pt(10)
                run.font.name = "Calibri"

    # ── Languages ────────────────────────────────────────────
    languages = optimized_resume.get("languages", [])
    if languages:
        _add_section_heading(doc, "Languages")
        lang_text = ", ".join(languages) if isinstance(languages, list) else str(languages)
        p = doc.add_paragraph()
        run = p.add_run(lang_text)
        run.font.size = Pt(10)
        run.font.name = "Calibri"

    # ── Write to buffer ──────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    logger.info("docx_generated", sections=len(doc.paragraphs))
    return buffer
